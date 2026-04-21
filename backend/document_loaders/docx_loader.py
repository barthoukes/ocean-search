#!/usr/bin/env python3
"""
Loader for documents.

This loader extracts basic metadata (dimensions, format, EXIF) from docx and text.
"""

import docx
from langchain_core.documents import Document

class DocxLoader(DocumentLoader):
    def _get_supported_extensions(self) -> Set[str]:
        return {'.docx'}
    
    def load_document(self, file_path: str) -> Optional[Document]:
        try:
            doc = docx.Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            metadata = {
                "filename": os.path.basename(file_path),
                "filepath": file_path,
                "type": "document",
                "extension": ".docx",
                "size": os.path.getsize(file_path),
                "match_source": "file_content",
                "embedding_type": "bert"  # or "ollama"
            }
            return Document(page_content=content, metadata=metadata)
        except Exception as e:
            print(f"Error loading DOCX: {e}")
            return None